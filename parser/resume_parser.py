"""
Resume Parser
=============

Pipeline:

    [PDF / DOCX / TXT bytes]
            │
            ▼
    extract_text()  →  pdfplumber / python-docx
            │
            ▼
    heuristic_parse()  →  regex + section detection
            │
            ▼
    llm_enrich()       →  GPT-4o-mini structured extraction (optional)
            │
            ▼
    ResumeStructure
"""

from __future__ import annotations

import io
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from utils.llm_client import chat_json, llm_available
from utils.text_cleaner import (
    clean_text,
    dedupe_keep_order,
    extract_emails,
    extract_phone_numbers,
    extract_urls,
    split_into_lines,
)


# --------------------------------------------------------------------------- #
# Data model
# --------------------------------------------------------------------------- #

@dataclass
class ResumeProject:
    name: str = ""
    description: str = ""
    technologies: List[str] = field(default_factory=list)


@dataclass
class ResumeExperience:
    company: str = ""
    title: str = ""
    duration: str = ""
    description: str = ""


@dataclass
class ResumeEducation:
    degree: str = ""
    institution: str = ""
    year: str = ""
    score: str = ""


@dataclass
class ResumeStructure:
    candidate_name: str = ""
    email: str = ""
    phone: str = ""
    links: List[str] = field(default_factory=list)
    location: str = ""

    summary: str = ""
    skills: List[str] = field(default_factory=list)
    tools: List[str] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)

    total_experience_years: float = 0.0
    experience: List[ResumeExperience] = field(default_factory=list)
    education: List[ResumeEducation] = field(default_factory=list)
    projects: List[ResumeProject] = field(default_factory=list)

    raw_text: str = ""
    source_filename: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)

    def search_corpus(self) -> str:
        """Concatenated string used for embedding similarity vs. JD."""
        proj_blob = " ".join(f"{p.name} {p.description} {' '.join(p.technologies)}" for p in self.projects)
        exp_blob = " ".join(f"{e.title} at {e.company} — {e.description}" for e in self.experience)
        return " | ".join(
            [self.summary, " ".join(self.skills), " ".join(self.tools), exp_blob, proj_blob]
        ).strip()


# --------------------------------------------------------------------------- #
# Text extraction
# --------------------------------------------------------------------------- #

def _extract_pdf(data: bytes) -> str:
    import pdfplumber

    pages: List[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(p for p in parts if p)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    if ext == ".docx":
        return _extract_docx(file_bytes)
    if ext in {".txt", ".md"}:
        return file_bytes.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported resume format: {ext}")


# --------------------------------------------------------------------------- #
# Heuristic parsing
# --------------------------------------------------------------------------- #

_NAME_HEADER_RE = re.compile(r"^(?:name)\s*[:\-]\s*([^\n]+)", re.IGNORECASE | re.MULTILINE)

_EXP_YEARS_RE = re.compile(
    r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)",
    re.IGNORECASE,
)

_SECTION_HEADS = {
    "summary":        ["summary", "profile", "objective", "about"],
    "skills":         ["skills", "technical skills", "key skills", "core skills"],
    "tools":          ["tools", "technologies", "tech stack"],
    "experience":     ["experience", "work experience", "professional experience", "employment"],
    "education":      ["education", "academics", "academic background"],
    "projects":       ["projects", "key projects", "personal projects", "academic projects"],
    "certifications": ["certifications", "courses", "certificates"],
}


def _section_blocks(text: str) -> Dict[str, str]:
    """Split text into named sections by heading detection."""
    lines = text.splitlines()

    # mark each line with its current section
    current = "preamble"
    blocks: Dict[str, List[str]] = {"preamble": []}

    for raw in lines:
        line = raw.strip()
        if not line:
            blocks.setdefault(current, []).append("")
            continue
        head_lower = line.lower().rstrip(":").rstrip()
        matched = None
        for section, aliases in _SECTION_HEADS.items():
            if any(head_lower == a or head_lower.startswith(a + " ") for a in aliases):
                matched = section
                break
        if matched:
            current = matched
            blocks.setdefault(current, [])
            continue
        blocks.setdefault(current, []).append(line)

    return {k: "\n".join(v).strip() for k, v in blocks.items()}


def _split_skills(block: str) -> List[str]:
    out: List[str] = []
    for line in split_into_lines(block):
        # remove leading bullet, then split on common separators
        cleaned = re.sub(r"^[•·\-*\s]+", "", line)
        # if there is a colon ("Languages: Python, SQL"), keep right side
        if ":" in cleaned:
            cleaned = cleaned.split(":", 1)[1]
        for piece in re.split(r"[,;|/]| and ", cleaned):
            p = piece.strip()
            if 1 <= len(p) <= 40:
                out.append(p)
    return dedupe_keep_order(out)


def _parse_education(block: str) -> List[ResumeEducation]:
    out: List[ResumeEducation] = []
    for line in split_into_lines(block):
        edu = ResumeEducation()
        year_match = re.search(r"(19|20)\d{2}", line)
        if year_match:
            edu.year = year_match.group(0)
        score_match = re.search(r"(?:CGPA|GPA|%)\s*[:\-]?\s*([\d.]+)", line, re.IGNORECASE)
        if score_match:
            edu.score = score_match.group(0)
        # crude split: "B.Tech in CS, IIT Madras, 2024"
        parts = [p.strip() for p in re.split(r"[,–—\-|]", line) if p.strip()]
        if parts:
            edu.degree = parts[0]
            if len(parts) > 1:
                edu.institution = parts[1]
        out.append(edu)
    return out


def _parse_experience(block: str) -> List[ResumeExperience]:
    out: List[ResumeExperience] = []
    chunks = re.split(r"\n\s*\n", block)
    for chunk in chunks:
        if not chunk.strip():
            continue
        lines = split_into_lines(chunk)
        if not lines:
            continue
        header = lines[0]
        exp = ResumeExperience()
        # try "Title at Company — Duration"
        m = re.match(r"(.+?)\s+(?:at|@)\s+(.+?)(?:\s*[—\-–]\s*(.+))?$", header, re.IGNORECASE)
        if m:
            exp.title = m.group(1).strip()
            exp.company = m.group(2).strip()
            exp.duration = (m.group(3) or "").strip()
        else:
            exp.title = header
        exp.description = " ".join(lines[1:])[:600]
        out.append(exp)
    return out


def _parse_projects(block: str) -> List[ResumeProject]:
    out: List[ResumeProject] = []
    chunks = re.split(r"\n\s*\n", block)
    for chunk in chunks:
        lines = split_into_lines(chunk)
        if not lines:
            continue
        proj = ResumeProject(name=lines[0].strip(" -*•·"))
        proj.description = " ".join(lines[1:])[:500]
        # naive tech extraction from "Tech: x, y, z" lines
        tech_match = re.search(r"(?:tech(?:nologies)?|stack|tools)\s*[:\-]\s*(.+)",
                               proj.description, re.IGNORECASE)
        if tech_match:
            proj.technologies = [t.strip() for t in re.split(r"[,;|/]", tech_match.group(1)) if t.strip()]
        out.append(proj)
    return out


def _guess_name(text: str) -> str:
    if m := _NAME_HEADER_RE.search(text):
        return m.group(1).strip()
    # fallback — first non-empty line that isn't an obvious header
    for line in split_into_lines(text)[:5]:
        if any(ch.isdigit() for ch in line):
            continue
        if "@" in line or "http" in line.lower():
            continue
        if 2 <= len(line.split()) <= 5 and len(line) < 60:
            return line
    return ""


def _heuristic_parse(text: str, filename: str) -> ResumeStructure:
    res = ResumeStructure(raw_text=text, source_filename=filename)
    flat = clean_text(text)

    res.candidate_name = _guess_name(text)
    emails = extract_emails(text)
    phones = extract_phone_numbers(text)
    res.email = emails[0] if emails else ""
    res.phone = phones[0] if phones else ""
    res.links = extract_urls(text)

    if m := _EXP_YEARS_RE.search(flat):
        try:
            res.total_experience_years = float(m.group(1))
        except ValueError:
            pass

    sections = _section_blocks(text)
    res.summary = sections.get("summary", "")[:800]
    res.skills = _split_skills(sections.get("skills", ""))
    res.tools = _split_skills(sections.get("tools", ""))
    res.certifications = [
        c.strip(" -*•·") for c in split_into_lines(sections.get("certifications", "")) if c.strip()
    ]
    res.experience = _parse_experience(sections.get("experience", ""))
    res.education = _parse_education(sections.get("education", ""))
    res.projects = _parse_projects(sections.get("projects", ""))

    return res


# --------------------------------------------------------------------------- #
# LLM enrichment
# --------------------------------------------------------------------------- #

_LLM_SYSTEM = """You are an expert resume parser. Extract STRICT JSON
representing the candidate's profile. Do not invent skills, experience or
projects. Use empty arrays / strings when unknown. Estimate
total_experience_years as a single float."""

_LLM_USER = """Extract this schema from the resume:

{{
  "candidate_name": str,
  "email": str,
  "phone": str,
  "location": str,
  "summary": str,
  "skills":         [str],
  "tools":          [str],
  "certifications": [str],
  "total_experience_years": float,
  "experience": [{{ "company": str, "title": str, "duration": str, "description": str }}],
  "education":  [{{ "degree": str, "institution": str, "year": str, "score": str }}],
  "projects":   [{{ "name": str, "description": str, "technologies": [str] }}]
}}

RESUME TEXT:
\"\"\"{resume}\"\"\"

Return JSON only."""


def _llm_enrich(text: str, base: ResumeStructure) -> ResumeStructure:
    if not llm_available():
        return base

    parsed = chat_json(_LLM_SYSTEM, _LLM_USER.format(resume=text[:9000]))
    if not parsed:
        return base

    base.candidate_name = parsed.get("candidate_name") or base.candidate_name
    base.email = parsed.get("email") or base.email
    base.phone = parsed.get("phone") or base.phone
    base.location = parsed.get("location") or base.location
    base.summary = parsed.get("summary") or base.summary

    if isinstance(parsed.get("skills"), list):
        base.skills = dedupe_keep_order([str(x).strip() for x in parsed["skills"] if str(x).strip()])
    if isinstance(parsed.get("tools"), list):
        base.tools = dedupe_keep_order([str(x).strip() for x in parsed["tools"] if str(x).strip()])
    if isinstance(parsed.get("certifications"), list):
        base.certifications = [str(x).strip() for x in parsed["certifications"] if str(x).strip()]

    try:
        years = float(parsed.get("total_experience_years") or base.total_experience_years)
        base.total_experience_years = years
    except (TypeError, ValueError):
        pass

    if isinstance(parsed.get("experience"), list):
        base.experience = [
            ResumeExperience(
                company=str(e.get("company", "")),
                title=str(e.get("title", "")),
                duration=str(e.get("duration", "")),
                description=str(e.get("description", "")),
            )
            for e in parsed["experience"] if isinstance(e, dict)
        ]

    if isinstance(parsed.get("education"), list):
        base.education = [
            ResumeEducation(
                degree=str(e.get("degree", "")),
                institution=str(e.get("institution", "")),
                year=str(e.get("year", "")),
                score=str(e.get("score", "")),
            )
            for e in parsed["education"] if isinstance(e, dict)
        ]

    if isinstance(parsed.get("projects"), list):
        base.projects = [
            ResumeProject(
                name=str(p.get("name", "")),
                description=str(p.get("description", "")),
                technologies=[str(t).strip() for t in (p.get("technologies") or []) if str(t).strip()],
            )
            for p in parsed["projects"] if isinstance(p, dict)
        ]

    return base


# --------------------------------------------------------------------------- #
# Public entry points
# --------------------------------------------------------------------------- #

def parse_resume_text(text: str, filename: str = "", use_llm: Optional[bool] = None) -> ResumeStructure:
    base = _heuristic_parse(text, filename)
    if use_llm is None:
        use_llm = llm_available()
    if use_llm:
        base = _llm_enrich(text, base)
    return base


def parse_resume(
    file: Union[bytes, str, Path],
    filename: Optional[str] = None,
    use_llm: Optional[bool] = None,
) -> ResumeStructure:
    """
    Parse a resume from either bytes (uploaded file) or a file path.
    """
    if isinstance(file, (str, Path)):
        path = Path(file)
        data = path.read_bytes()
        fname = filename or path.name
    else:
        data = file
        fname = filename or "resume.pdf"

    text = extract_text(data, fname)
    return parse_resume_text(text, filename=fname, use_llm=use_llm)
